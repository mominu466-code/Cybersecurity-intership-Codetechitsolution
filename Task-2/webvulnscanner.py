import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse, urljoin
from datetime import datetime
import time
import os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from colorama import Fore, Style, init

init(autoreset=True)

# ================= CONFIG =================

DEFAULT_PARAM = "test"

SQL_PAYLOADS_BASIC = [
    "' OR '1'='1",
    "' OR 1=1 --",
    "\" OR \"1\"=\"1"
]

XSS_PAYLOADS_BASIC = [
    "<script>alert(1)</script>",
    "\"><script>alert(1)</script>",
    "<img src=x onerror=alert(1)>"
]

SQL_PAYLOADS_ADV = [
    "' OR '1'='1",
    "' OR 1=1 --",
    "' OR 'a'='a' --",
    "' AND 1=2 --",
    "' UNION SELECT NULL --"
]

XSS_PAYLOADS_ADV = [
    "<script>alert(1)</script>",
    "\"><script>alert(1)</script>",
    "<img src=x onerror=alert(1)>",
    "<svg/onload=alert(1)>"
]

SQL_ERRORS = [
    "you have an error in your sql syntax",
    "warning: mysql",
    "sql error",
    "unclosed quotation mark"
]

# ================= STATE =================

findings = []
current_target = None
last_scan = {"mode": None, "duration": 0, "note": None}

# ================= HELPERS =================

def request(url, **kw):
    try:
        return requests.get(url, timeout=5, **kw)
    except requests.RequestException:
        return None


def record(vtype, component, payload, mode, poc):
    severity = "High" if vtype == "SQL Injection" else "Medium"
    fix = (
        "Use parameterized queries and validate user input."
        if vtype == "SQL Injection"
        else "Sanitize input and encode output."
    )

    findings.append({
        "type": vtype,
        "component": component,
        "payload": payload,
        "mode": mode,
        "poc": poc,
        "severity": severity,
        "fix": fix
    })

# ================= SCANNERS =================

def scan_sql(url, payloads, mode):
    print(Fore.CYAN + "\n[+] SQL Injection Scan Started")
    start = time.time()

    parsed = urlparse(url)
    params = parse_qs(parsed.query) or {DEFAULT_PARAM: ["1"]}
    hit = False

    for p in params:
        for payload in payloads:
            print(Fore.WHITE + f"    Testing SQL payload → {payload}")
            q = params.copy()
            q[p] = payload
            test_url = urlunparse(parsed._replace(query=urlencode(q, doseq=True)))

            r = request(test_url)
            if not r:
                continue

            for err in SQL_ERRORS:
                if err in r.text.lower():
                    hit = True
                    print(Fore.RED + f"    [!] SQL Injection FOUND → payload: {payload}")
                    record("SQL Injection", f"URL parameter '{p}'", payload, mode, test_url)
                    break

    last_scan["duration"] = round(time.time() - start, 2)
    if not hit:
        print(Fore.GREEN + "[-] SQL Injection NOT FOUND")
        last_scan["note"] = "No SQL Injection detected"


def scan_xss(url, payloads, mode):
    print(Fore.CYAN + "\n[+] XSS Scan Started")
    start = time.time()

    parsed = urlparse(url)
    params = parse_qs(parsed.query)
    hit = False

    # URL-based XSS
    if not params:
        for payload in payloads:
            print(Fore.WHITE + f"    Testing XSS payload → {payload}")
            test_url = f"{url}?{DEFAULT_PARAM}={payload}"
            r = request(test_url)

            if r and payload in r.text:
                hit = True
                print(Fore.RED + f"    [!] XSS FOUND → payload: {payload}")
                record(
                    "Cross-Site Scripting (XSS)",
                    f"URL parameter '{DEFAULT_PARAM}'",
                    payload,
                    mode,
                    test_url
                )

        last_scan["duration"] = round(time.time() - start, 2)
        if not hit:
            print(Fore.GREEN + "[-] XSS NOT FOUND")
            last_scan["note"] = "No XSS detected"
        return

    # Form-based XSS
    r = request(url)
    if not r:
        last_scan["note"] = "Connection error"
        return

    soup = BeautifulSoup(r.text, "html.parser")
    forms = soup.find_all("form")

    for form in forms:
        action = urljoin(url, form.get("action", ""))
        inputs = form.find_all("input")

        for payload in payloads:
            print(Fore.WHITE + f"    Testing XSS payload → {payload}")
            data = {i.get("name"): payload for i in inputs if i.get("name")}
            r2 = request(action, params=data)

            if r2 and payload in r2.text:
                hit = True
                print(Fore.RED + f"    [!] XSS FOUND → payload: {payload}")
                record(
                    "Cross-Site Scripting (XSS)",
                    f"Form action {action}",
                    payload,
                    mode,
                    action
                )

    last_scan["duration"] = round(time.time() - start, 2)
    if not hit:
        print(Fore.GREEN + "[-] XSS NOT FOUND")
        last_scan["note"] = "No XSS detected"

# ================= MODES =================

def basic_scan():
    last_scan.update({"mode": "Basic", "note": None})
    scan_sql(current_target, SQL_PAYLOADS_BASIC, "Basic")
    scan_xss(current_target, XSS_PAYLOADS_BASIC, "Basic")


def advanced_scan():
    last_scan.update({"mode": "Advanced", "note": None})
    scan_sql(current_target, SQL_PAYLOADS_ADV, "Advanced")
    scan_xss(current_target, XSS_PAYLOADS_ADV, "Advanced")

# ================= REPORTS =================

def generate_docx():
    os.makedirs("data", exist_ok=True)
    name = input("Report name: ").strip()
    path = f"data/{name}.docx"

    doc = Document()
    doc.add_heading("Web Vulnerability Assessment Report", 1)
    doc.add_paragraph(f"Target URL: {current_target}")
    doc.add_paragraph(f"Scan Time : {datetime.now()}")

    if not findings:
        doc.add_heading("Result: No Vulnerability Found", 2)
        doc.save(path)
        print(Fore.GREEN + f"[+] Report saved → {path}")
        return

    doc.add_heading("Result: Vulnerabilities Found", 2)

    for i, f in enumerate(findings, 1):
        doc.add_heading(f"Finding {i}", 3)
        doc.add_paragraph(f"Type: {f['type']}")
        doc.add_paragraph(f"Severity: {f['severity']}")
        doc.add_paragraph(f"Component: {f['component']}")
        doc.add_paragraph(f"Payload: {f['payload']}")
        doc.add_paragraph(f"PoC: {f['poc']}")
        doc.add_paragraph(f"Recommendation: {f['fix']}")

    doc.save(path)
    print(Fore.GREEN + f"[+] Report saved → {path}")


def generate_pdf():
    os.makedirs("data", exist_ok=True)
    name = input("Report name: ").strip()
    path = f"data/{name}.pdf"

    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    y = h - 50

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Web Vulnerability Assessment Report")
    y -= 30

    c.setFont("Helvetica", 11)
    c.drawString(50, y, f"Target URL: {current_target}")
    y -= 15
    c.drawString(50, y, f"Scan Time : {datetime.now()}")
    y -= 30

    if not findings:
        c.drawString(50, y, "Result: No Vulnerability Found")
        c.save()
        print(Fore.GREEN + f"[+] Report saved → {path}")
        return

    c.drawString(50, y, "Result: Vulnerabilities Found")
    y -= 25

    for i, f in enumerate(findings, 1):
        c.drawString(50, y, f"Finding {i}")
        y -= 15
        c.drawString(60, y, f"Type: {f['type']}"); y -= 12
        c.drawString(60, y, f"Severity: {f['severity']}"); y -= 12
        c.drawString(60, y, f"Component: {f['component']}"); y -= 12
        c.drawString(60, y, f"Payload: {f['payload']}"); y -= 12
        c.drawString(60, y, f"PoC: {f['poc']}"); y -= 12
        c.drawString(60, y, f"Recommendation: {f['fix']}"); y -= 20

        if y < 100:
            c.showPage()
            y = h - 50

    c.save()
    print(Fore.GREEN + f"[+] Report saved → {path}")


def choose_report_format():
    print(Fore.CYAN + "\nSelect Report Format")
    print(Fore.CYAN + "1. DOCX (Word)")
    print(Fore.CYAN + "2. PDF")

    ch = input("Select: ").strip()
    if ch == "1":
        generate_docx()
    elif ch == "2":
        generate_pdf()
    else:
        print(Fore.RED + "Invalid selection")

# ================= DASHBOARD =================

def dashboard():
    global current_target

    while True:
        print(Style.BRIGHT + Fore.BLUE + "\n" + "=" * 50)
        print(Fore.BLUE + "   WEB APPLICATION VULNERABILITY SCANNER")
        print(Fore.BLUE + "=" * 50)

        print(Fore.WHITE + f"Target URL : {current_target}")
        print(Fore.WHITE + f"Last Scan  : {last_scan['mode']}")
        print(Fore.GREEN + f"Findings   : {len(findings)}")
        if last_scan["note"]:
            print(Fore.YELLOW + f"Note       : {last_scan['note']}")

        print(Fore.BLUE + "-" * 50)
        print(Fore.CYAN + "1. Basic Scan (SQL + XSS)")
        print(Fore.CYAN + "2. Advanced Scan (Deep Scan)")
        print(Fore.CYAN + "3. Generate Report")
        print(Fore.CYAN + "4. Change Target URL")
        print(Fore.CYAN + "5. Clear Findings")
        print(Fore.CYAN + "6. Exit")
        print(Fore.BLUE + "-" * 50)

        choice = input(Fore.WHITE + "Select option: ").strip()

        if choice == "1":
            basic_scan()
        elif choice == "2":
            advanced_scan()
        elif choice == "3":
            if not findings:
                print(Fore.YELLOW + "[-] No findings to report")
            else:
                choose_report_format()
        elif choice == "4":
            findings.clear()
            last_scan.update({"mode": None, "duration": 0, "note": None})
            current_target = input("Enter new target URL: ").strip()
            print(Fore.GREEN + "[+] Target updated")
        elif choice == "5":
            findings.clear()
            last_scan.update({"mode": None, "duration": 0, "note": None})
            print(Fore.GREEN + "[+] Findings cleared")
        elif choice == "6":
            print(Fore.GREEN + "[+] Exiting tool")
            break
        else:
            print(Fore.RED + "Invalid option")

# ================= MAIN =================

if __name__ == "__main__":
    print(Style.BRIGHT + Fore.BLUE + "\n=========================================")
    print(Fore.BLUE + "   WEB APPLICATION VULNERABILITY SCANNER")
    print(Style.BRIGHT + Fore.BLUE + "=========================================")
    current_target = input("Enter Target URL: ").strip()
    dashboard()
